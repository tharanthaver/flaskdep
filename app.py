import os
import uuid
import logging
import requests
import io
import contextlib
import random
import time
from flask import Flask, request, jsonify, send_file, render_template
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont

# Flask app initialization
app = Flask(__name__, template_folder='templates')

UPLOAD_FOLDER = 'solved_files'
TEMP_FOLDER = 'temp'

# Create necessary folders
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(TEMP_FOLDER, exist_ok=True)

# API Key (replace with actual key)
API_KEY = os.getenv("API_KEY")
BASE_URL = "https://api.hyperbolic.xyz/v1/chat/completions"

# Logging configuration
logging.basicConfig(level=logging.DEBUG,
                    format='%(asctime)s - %(levelname)s - %(message)s',
                    handlers=[logging.StreamHandler(), logging.FileHandler('app.log')])

# ----- Utility Functions -----
def extract_text_from_pdf(pdf_path):
    try:
        with open(pdf_path, "rb") as pdf_file:
            reader = PdfReader(pdf_file)
            text = "\n".join(page.extract_text() for page in reader.pages if page.extract_text()).strip()
            return text.strip()
    except Exception as e:
        logging.exception(f"Error extracting text from PDF: {e}")
        return f"Error extracting text from PDF: {e}"


def solve_coding_problem(question):
    try:
        # Detect if the question expects user input and replace it with a fixed random number
        if "input" in question.lower() or "user" in question.lower():
            random_number = random.randint(1, 7)
            question = f"{question} Assume the user input is {random_number}. The code should NOT prompt for input."

        prompt = f"""
        Write only the Python code to solve the following problem: {question}.
        The code must NOT contain any 'input()' function or request user input.
        Instead, assume predefined values for any required inputs.
        Do not include any explanations, comments, or unnecessary imports. Just provide the clean Python code.
        """

        headers = {"Content-Type": "application/json", "Authorization": f"Bearer {API_KEY}"}
        data = {"messages": [{"role": "user", "content": prompt}], "model": "meta-llama/Llama-3.3-70B-Instruct", "max_tokens": 512}

        for attempt in range(5):
            response = requests.post(BASE_URL, headers=headers, json=data)
            
            if response.status_code == 200:
                response_json = response.json()
                solution = response_json.get("choices", [{}])[0].get("message", {}).get("content", "").strip()

                # Clean the solution
                solution = solution.replace("```python", "").replace("```", "").strip()

                if solution and "input(" not in solution:  # Ensure the response does not ask for input
                    return solution
                else:
                    logging.warning("Generated code still contains 'input()'. Retrying...")
            elif response.status_code == 429:
                retry_after = int(response.headers.get('Retry-After', 1))
                logging.warning(f"Rate limit hit. Retrying after {retry_after} seconds.")
                time.sleep(retry_after)
            else:
                logging.error(f"API error: {response.status_code} - {response.text}")
                break

        return "Error: Unable to fetch solution from the API after multiple attempts."
    except Exception as e:
        logging.error(f"Error solving coding problem: {e}")
        return f"Error solving coding problem: {e}"


def execute_code(code):
    try:
        if code.startswith("Error") or not code:
            return f"Invalid code returned: {code}"

        logging.debug(f"Executing the following code:\n{code}")

        output = io.StringIO()
        context = {}

        with contextlib.redirect_stdout(output), contextlib.redirect_stderr(output):
            exec(code, context, context)

        result = output.getvalue().strip()
        return result if result else "Error: Code execution produced no output."

    except Exception as e:
        logging.exception(f"Error executing code: {e}")
        return f"Error executing code: {e}"

def create_screenshot(output_text):
    try:
        output_text = output_text.strip() or "No output."
        font_path = "C:\\Windows\\Fonts\\consola.ttf"
        font = ImageFont.truetype(font_path, 14)
        lines = output_text.splitlines()
        width, height = 800, max(len(lines) * 16 + 20, 100)
        
        image = Image.new('RGB', (width, height), color='black')
        draw = ImageDraw.Draw(image)
        y_offset = 10
        
        for line in lines:
            draw.text((10, y_offset), line, fill='white', font=font)
            y_offset += 16
        
        buffer = io.BytesIO()
        image.save(buffer, format="PNG")
        buffer.seek(0)
        return buffer.read()
    except Exception as e:
        logging.exception(f"Error creating screenshot: {e}")
        return b""


def generate_word_doc(name, reg_number, questions, solutions, screenshots, output_path):
    """
    Creates a Word document with a well-formatted structure:
    - Question
    - Code Solution
    - Screenshot
    """
    try:
        doc = Document()
        doc.add_heading(f"Name: {name}", level=1)
        doc.add_heading(f"Register Number: {reg_number}", level=1)
        doc.add_paragraph("\n")

        for i, (question, solution, screenshot) in enumerate(zip(questions, solutions, screenshots)):
            doc.add_heading(f"Question {i + 1}:", level=2)
            doc.add_paragraph(question.strip())

            doc.add_heading("Code Solution:", level=3)
            doc.add_paragraph(solution.strip())

            doc.add_heading("Final Output Screenshot:", level=3)
            if screenshot:
                doc.add_picture(io.BytesIO(screenshot), width=Inches(6))

            doc.add_page_break()

        doc.save(output_path)
        return output_path
    except Exception as e:
        logging.exception(f"Error generating Word document: {e}")
        return f"Error generating Word document: {e}"


# ----- Routes -----
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload_pdf', methods=['POST'])
def upload_pdf():
    try:
        file = request.files['file']
        name = request.form.get('name')
        reg_number = request.form.get('regNo')

        if not file or not name or not reg_number:
            return jsonify({"error": "Name, Register Number, and PDF file are required."}), 400

        filename = secure_filename(file.filename)
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(file_path)

        pdf_text = extract_text_from_pdf(file_path)
        questions = pdf_text.split('\n')
        solutions = [solve_coding_problem(q) for q in questions]
        final_outputs = [execute_code(sol) for sol in solutions]
        screenshots = [create_screenshot(output) for output in final_outputs]

        output_file = os.path.join(TEMP_FOLDER, f"solutions_{uuid.uuid4().hex}.docx")
        generate_word_doc(name, reg_number, questions, solutions, screenshots, output_file)

        return send_file(output_file, as_attachment=True)
    except Exception as e:
        logging.exception(f"Error processing PDF: {e}")
        return jsonify({'error': str(e)}), 500



import re

def split_questions(text):
    """
    Splits pasted text into individual coding questions based on numbering and common patterns.
    """
    numbered_split = re.split(r'\n\s*\d+\.\s*', text.strip())

    if len(numbered_split) > 1:
        return [q.strip() for q in numbered_split if q.strip()]

    question_patterns = [
        r'(?i)(?=\bWrite a program\b)',
        r'(?i)(?=\bSolve this problem\b)',
        r'(?i)(?=\bImplement a function\b)'
    ]
    split_text = re.split('|'.join(question_patterns), text)

    return [q.strip() for q in split_text if q.strip()]

@app.route('/manual_solve', methods=['POST'])
def manual_solve():
    try:
        data = request.get_json()
        name = data.get('name')
        reg_number = data.get('regNo')
        questions_text = data.get('questions')

        if not questions_text or not name or not reg_number:
            return jsonify({"error": "Name, Register Number, and Questions are required."}), 400

        # Auto-split questions if they are in a single block
        questions = split_questions(questions_text) if isinstance(questions_text, str) else questions_text

        solutions = [solve_coding_problem(q) for q in questions]
        final_outputs = [execute_code(sol) for sol in solutions]
        screenshots = [create_screenshot(output) for output in final_outputs]

        output_file = os.path.join(TEMP_FOLDER, f"manual_solutions_{uuid.uuid4().hex}.docx")
        generated_path = generate_word_doc(name, reg_number, questions, solutions, screenshots, output_file)

        if not os.path.exists(generated_path):
            return jsonify({"error": "Failed to generate the document."}), 500

        return send_file(generated_path, as_attachment=True)

    except Exception as e:
        logging.exception(f"Error processing manual questions: {e}")
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True)
